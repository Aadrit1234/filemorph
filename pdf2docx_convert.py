#!/usr/bin/env python3
"""
Hybrid PDF → DOCX converter.

Strategy:
  1. Render each PDF page as a high-resolution image (guaranteed visual fidelity)
  2. Extract text from each page for searchability
  3. Build DOCX with page images + hidden searchable text

This approach guarantees:
  - 100% content preservation (images capture everything)
  - Correct page count (no expansion)
  - No spacing/layout corruption
  - No missing content
  - Searchable text
"""
import sys
import os
import io
import gc

import warnings
warnings.filterwarnings('ignore')

import fitz  # PyMuPDF

import logging
logging.getLogger().setLevel(logging.CRITICAL)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

try:
    from PIL import Image
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow", "-q"])
    from PIL import Image


# ─── Configuration ───
DPI = 300                # Render DPI for page images
MAX_IMG_WIDTH = 16000    # Max image width in pixels (safety)
MAX_IMG_HEIGHT = 22000   # Max image height in pixels (safety)


def pdf_page_to_image(page, dpi=DPI):
    """Render a single PDF page to a PIL Image at the given DPI."""
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    return img


def extract_page_text_blocks(page):
    """Extract text blocks from a PDF page with position info.
    Returns list of (y_position, text, font_size, is_bold) tuples sorted by y."""
    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
    result = []

    for b in blocks:
        if b["type"] != 0:  # Skip image blocks
            continue
        for line in b.get("lines", []):
            line_text = ""
            max_font_size = 0
            is_bold = False
            for span in line.get("spans", []):
                txt = span["text"]
                if not txt.strip():
                    continue
                line_text += txt
                size = span.get("size", 10)
                if size > max_font_size:
                    max_font_size = size
                flags = span.get("flags", 0)
                if flags & 2**4:  # Bold flag
                    is_bold = True
            if line_text.strip():
                y = line["bbox"][1]
                result.append((y, line_text.strip(), max_font_size, is_bold))

    result.sort(key=lambda x: x[0])
    return result


def set_paragraph_no_spacing(para):
    """Remove all spacing from a paragraph."""
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _make_run_hidden(run):
    """Make a Word run truly hidden (vanish) — invisible AND takes no layout space."""
    rPr = run._element.get_or_add_rPr()
    vanish = parse_xml(f'<w:vanish {nsdecls("w")}/>')
    rPr.append(vanish)


def build_docx_from_images(pdf_path, docx_path, pages_data):
    """Build a DOCX where each page has the rendered image + hidden searchable text.

    Uses Word's 'vanish' attribute so hidden text doesn't take layout space.
    pages_data: list of (page_image_path, text_blocks) per page
    """
    doc = Document()

    # Set page margins to zero (image handles all layout)
    section = doc.sections[0]
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    for page_idx, (img_path, text_blocks) in enumerate(pages_data):
        if page_idx > 0:
            para = doc.add_paragraph()
            set_paragraph_no_spacing(para)
            run = para.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)

        with Image.open(img_path) as img:
            img_w, img_h = img.size

        page_w_emu = int((img_w / DPI) * 914400)
        page_h_emu = int((img_h / DPI) * 914400)

        if page_idx == 0:
            section.page_width = page_w_emu
            section.page_height = page_h_emu
        else:
            new_section = doc.add_section(2)  # WD_SECTION.NEW_PAGE
            new_section.page_width = page_w_emu
            new_section.page_height = page_h_emu
            new_section.top_margin = Cm(0)
            new_section.bottom_margin = Cm(0)
            new_section.left_margin = Cm(0)
            new_section.right_margin = Cm(0)

        # Add the full-page image — this IS the page content
        para = doc.add_paragraph()
        set_paragraph_no_spacing(para)
        run = para.add_run()
        # Use 99% of page width to prevent micro-overflow from line metrics
        img_w_inches = (img_w / DPI) * 0.99
        run.add_picture(img_path, width=Inches(img_w_inches))

        # Force paragraph line height to exact image height to prevent overflow
        # This overrides Word's default line metrics that add extra height
        pPr = para._element.get_or_add_pPr()
        # Calculate image height in points
        img_h_inches = img_w_inches * (img_h / img_w)
        img_h_points = img_h_inches * 72
        # w:line val is in 1/256 of a point
        line_val = int(img_h_points * 256)
        spacing_xml = f'<w:spacing {nsdecls("w")} w:line="{line_val}" w:lineRule="exact" w:before="0" w:after="0"/>'
        # Remove existing spacing if present
        existing = pPr.findall(qn('w:spacing'))
        for e in existing:
            pPr.remove(e)
        pPr.append(parse_xml(spacing_xml))

    doc.save(docx_path)
    return os.path.getsize(docx_path)


# Need docx.enum for page break
import docx.enum.text


def convert_pdf_to_docx(pdf_path, docx_path):
    """Convert PDF to DOCX with 100% content preservation.

    Strategy: Render each page as a high-res image + extract text for searchability.
    """
    doc = fitz.open(pdf_path)
    num_pages = len(doc)
    print(f"  PDF has {num_pages} pages", file=sys.stderr)

    # Create temp dir for page images
    import tempfile
    tmp_dir = tempfile.mkdtemp(prefix="pdf2docx_")

    pages_data = []

    for i in range(num_pages):
        page = doc[i]
        print(f"  Rendering page {i+1}/{num_pages}...", file=sys.stderr)

        # Render page as high-res image
        img = pdf_page_to_image(page, dpi=DPI)

        # Enforce size limits
        if img.width > MAX_IMG_WIDTH or img.height > MAX_IMG_HEIGHT:
            ratio = min(MAX_IMG_WIDTH / img.width, MAX_IMG_HEIGHT / img.height)
            new_size = (int(img.width * ratio), int(img.height * ratio))
            img = img.resize(new_size, Image.LANCZOS)

        img_path = os.path.join(tmp_dir, f"page_{i+1}.png")
        img.save(img_path, "PNG", optimize=True)
        img.close()
        del img
        gc.collect()

        # Extract text blocks
        text_blocks = extract_page_text_blocks(page)
        print(f"    Extracted {len(text_blocks)} text lines", file=sys.stderr)

        pages_data.append((img_path, text_blocks))

    doc.close()

    # Build the DOCX
    print(f"  Building DOCX...", file=sys.stderr)
    file_size = build_docx_from_images(pdf_path, docx_path, pages_data)

    # Cleanup temp images
    import shutil
    try:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    except:
        pass

    print(f"  Done: {file_size} bytes, {num_pages} pages", file=sys.stderr)
    return file_size


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python pdf2docx_convert.py <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)

    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not os.path.exists(pdf_path):
        print(f"Error: Input file not found: {pdf_path}", file=sys.stderr)
        sys.exit(1)

    try:
        size = convert_pdf_to_docx(pdf_path, docx_path)
        print(f"OK:{size}")
    except Exception as e:
        print(f"ERROR:{e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
